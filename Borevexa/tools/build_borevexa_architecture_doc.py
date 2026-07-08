from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(__file__).resolve().parents[1] / "Borevexa_Prescan_App_Werking_Layout_Structuur.docx"


BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "071422"
MUTED = "587080"
LIGHT_FILL = "F2F4F7"
CALLOUT = "F4F6F9"
GREEN_FILL = "E6F5ED"
YELLOW_FILL = "FFF7E6"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_cell_margins(table, top=80, bottom=80, start=120, end=120):
    tbl_pr = table._tbl.tblPr
    margins = tbl_pr.find(qn("w:tblCellMar"))
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    values = {"top": top, "bottom": bottom, "start": start, "end": end}
    for key, value in values.items():
        node = margins.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def style_run(run, bold=None, color=None, size=None):
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if size:
        run.font.size = Pt(size)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    return p


def add_body(doc, text):
    p = doc.add_paragraph(style="Normal")
    p.add_run(text)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_callout(doc, title, body, fill=CALLOUT):
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [9360])
    set_cell_margins(table, top=120, bottom=120, start=160, end=160)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    r = p.add_run(title)
    style_run(r, bold=True, color=DARK_BLUE)
    p2 = cell.add_paragraph()
    p2.add_run(body)
    doc.add_paragraph()
    return table


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr[i].text = header
        set_cell_shading(hdr[i], LIGHT_FILL)
        for paragraph in hdr[i].paragraphs:
            for run in paragraph.runs:
                style_run(run, bold=True, color=INK)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    set_table_width(table, widths)
    set_cell_margins(table)
    doc.add_paragraph()
    return table


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    bullet = styles["List Bullet"]
    bullet.font.name = "Calibri"
    bullet.font.size = Pt(11)
    bullet.paragraph_format.left_indent = Inches(0.5)
    bullet.paragraph_format.first_line_indent = Inches(-0.25)
    bullet.paragraph_format.space_after = Pt(8)
    bullet.paragraph_format.line_spacing = 1.167

    footer = section.footer.paragraphs[0]
    footer.text = "Borevexa Prescan Native - technische documentatie"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor.from_string(MUTED)


def build_doc():
    doc = Document()
    configure_document(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = title.add_run("Borevexa Prescan Native")
    style_run(r, bold=True, color=INK, size=24)
    subtitle = doc.add_paragraph()
    r = subtitle.add_run("Werking, layout, structuur en fase 1 t/m 4")
    style_run(r, color=MUTED, size=12)
    meta = doc.add_paragraph()
    meta.add_run("Versiecontext: v322 | Datum: 16-06-2026 | Doel: overdracht en doorontwikkeling")

    add_callout(
        doc,
        "Kernbeeld",
        "De app is omgebouwd van losse stapinhoud naar een rapportgedreven workflow: processtap invullen, substapdata opslaan, rapportcontract controleren, preview bekijken en exporteren met manifest en historie.",
        GREEN_FILL,
    )

    add_heading(doc, "1. Doel van de app", 1)
    add_body(
        doc,
        "Borevexa Prescan Native is een WPF desktopapp voor de voorbereiding en rapportage van HDD/prescan projecten. De app combineert projectinformatie, importbestanden, kaartlagen, boorlijn, BGT/KLIC/BAG-context, dwarsprofiel, machinekeuze en eindrapportage in een procesgestuurde workflow."
    )
    add_body(
        doc,
        "Het huidige doel is dat de gebruiker geen losse rapportteksten of plakplaatjes meer hoeft te maken. Elke substap levert eigen rapportdata op; de eindrapportage wordt automatisch opgebouwd vanuit die data."
    )

    add_heading(doc, "2. Hoofdstructuur van de workflow", 1)
    add_table(
        doc,
        ["Stap", "Doel", "Belangrijkste output"],
        [
            ["0", "Voorblad, voorwoord en inhoudsopgave", "Rapportfrontmatter en inhoudsstructuur"],
            ["1", "Projectinformatie", "Projectgegevens, inhoud, vulgraad, machinekeuze"],
            ["2", "Ontwerp, KLIC, BAG en BGT inladen", "Importbestanden, documenten, lagen en kruisingen"],
            ["3", "Boorlijn", "Boorlijn ingetekend en KLIC kruisingen"],
            ["4", "Oppervlakteanalyse", "BGT segmenten en analyse langs boorlijn"],
            ["5", "Omgevingsmanagement", "Perceelsegmenten en ZRO analyse"],
            ["6", "Ondergrondanalyse", "BRO/AHN bronnen"],
            ["7", "Dwarsprofiel", "Dwarsprofiel ingetekend"],
            ["8", "Machine locatie", "Machine ingetekend"],
            ["9", "Sonderingen", "Sonderingen ingetekend"],
            ["10", "Eindrapport en export", "Rapportpreview, statusdashboard, exportmanifest"],
            ["11", "3D/profiel vervolg", "3D context en export"],
            ["12", "Werktekening", "Werktekening en diagnose"],
        ],
        [900, 3300, 5160],
    )

    add_heading(doc, "3. Layout en bediening", 1)
    add_bullets(
        doc,
        [
            "Topbar: bevat Projecten, Bestanden, Vorige, Volgende, Rapportpreview en centrale projectopslag.",
            "Processtappenbalk links: toont hoofdstappen en substappen; ingeklapt blijven nummers zichtbaar en klikbaar.",
            "Werkgebied midden: toont de actieve substap, bijvoorbeeld projectgegevens, inhoud, kaart, boorlijn of analyse.",
            "Rechter zijbalk: bevat contexttools zoals kaartlagen of boorlijnbediening, afhankelijk van de stap.",
            "Rapportpreview rechts: uitschuifbare viewer met zoom en export; toont het rapportonderdeel van de actieve stap/substap.",
        ],
    )

    add_heading(doc, "4. Rapportgedreven architectuur", 1)
    add_table(
        doc,
        ["Onderdeel", "Verantwoordelijkheid"],
        [
            ["StepReportCatalog", "Definieert welke substappen onder een hoofdstap vallen."],
            ["ReportContractCatalog", "Legt per substap vast welke bronnen vereist of optioneel zijn."],
            ["ReportQualityService", "Bepaalt rapportstatus: Niet gestart, Onvolledig, Controle nodig of Rapportklaar."],
            ["ReportRenderService", "Routeert substapdata naar nette rapportblokken en renderers."],
            ["ReportPreviewService", "Beheert live kaartcaptures, kaartlocks en previewbeelden."],
            ["ReportExportService", "Regelt exportbestanden, manifesten, status, versie en historie."],
            ["MapStateService", "Slaat kaartpositie, schaal, lagen en kaartcontext op."],
            ["ProjectRepository", "Lokale opslag van projectdata, stepdata, importdata en exportmetadata."],
        ],
        [2300, 7060],
    )

    add_heading(doc, "5. Rapportdata per substap", 1)
    add_body(
        doc,
        "Elke substap schrijft eigen rapportdata weg onder step_report_data. De hoofdstap is vooral navigatie en container; inhoud hoort in de substappen. Hierdoor kan de eindrapportage automatisch per onderdeel worden samengesteld en gecontroleerd."
    )
    add_callout(
        doc,
        "Belangrijk principe",
        "Een substap is pas rapportwaardig als de contractbronnen aanwezig zijn en de renderer een leesbare rapportsectie kan tonen. Ruwe JSON wordt niet als eindrapportinhoud gebruikt.",
        CALLOUT,
    )

    add_heading(doc, "6. Fases en huidige status", 1)
    add_table(
        doc,
        ["Fase", "Status", "Resultaat"],
        [
            ["Fase 1", "Gereed", "Stap/substapstructuur, scrollbare sidebar, hoofd- en substapnavigatie, gecentraliseerde opslag."],
            ["Fase 2", "Gereed", "Substaprenderers, leesbare rapportpreview, kaart- en doorsnedeweergaves zonder ruwe JSON."],
            ["Fase 3", "Gereed voor deze ronde", "Rapportcontracten, kwaliteitsstatus, dashboard en exportblokkade bij hoge issues."],
            ["Fase 4", "Gereed voor deze ronde", "Exportservice, concept/definitief, manifesten, exportgeschiedenis en metadata."],
        ],
        [1300, 2100, 5960],
    )

    add_heading(doc, "7. Kaart en rapportpreview", 1)
    add_body(
        doc,
        "Voor kaartstappen gebruikt de rapportage het kaartbeeld uit de processtap. De zichtbare kaart, laagkeuze, zoom, schaal en boorlijncontext worden vastgelegd of live gecaptured. De rapportpreview moet overeenkomen met wat de gebruiker in de processtap ziet, zonder UI-knoppen zoals zoom, hand, lijn of wis."
    )
    add_bullets(
        doc,
        [
            "Kaartlagen voor rapportage zijn bewust beperkt tot relevante PDOK lagen, overlays en boorlijnelementen.",
            "Rapportkaartbeelden worden als afbeelding gebruikt om renderproblemen met tiles te vermijden.",
            "Bij export wordt een manifest gemaakt zodat herleidbaar is welke versie, status en bestanden zijn gebruikt.",
        ],
    )

    add_heading(doc, "8. Export en versiebeheer", 1)
    add_table(
        doc,
        ["Exporttype", "Bestanden", "Registratie"],
        [
            ["Eindrapport", "HTML, Markdown en manifest JSON", "Opgeslagen onder eindrapport_export en eindrapport_export_history."],
            ["Rapportpreview", "HTML, PNG en manifest JSON", "Opgeslagen per stap/substap met preview-exportgeschiedenis."],
            ["Status", "CONCEPT of DEF", "Bepaald door ReportQualitySummary.IsReady."],
        ],
        [1900, 3500, 3960],
    )

    add_heading(doc, "9. Wat nog later kan worden verbeterd", 1)
    add_bullets(
        doc,
        [
            "Native PDF-rendering zonder browser/print-route.",
            "Meer inhoudelijke validatie per substap, bijvoorbeeld kaart gevuld, boorlijn geldig, kruisingen gecontroleerd.",
            "Automatische regressietests voor ReportQualityService, ReportExportService en belangrijke renderers.",
            "Nog meer layoutpolijsting van het eindrapport, met vaste hoofdstuktemplates en printprofielen.",
            "Een projectbreed rapportcontrolepaneel in de topbar of sidebar.",
        ],
    )

    add_heading(doc, "10. Praktisch gebruik", 1)
    add_bullets(
        doc,
        [
            "Vul een substap in en sla centraal het project op.",
            "Controleer de rapportpreview aan de rechterkant.",
            "Gebruik de rapportstatus om te zien wat nog mist.",
            "Gebruik stap 10 voor het eindrapport, exportgeschiedenis en definitieve export.",
            "Gebruik de manifestbestanden voor herleidbaarheid bij oplevering.",
        ],
    )

    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(build_doc())
